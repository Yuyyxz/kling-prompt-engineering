from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── 页面设置 ──
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

# ── 样式设置 ──
style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# Helper functions
def add_heading_styled(text, level=1, color=RGBColor(0x2B, 0x3A, 0x4E)):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = color
        run.font.name = 'Microsoft YaHei'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return h

def add_para(text, bold=False, italic=False, size=11, color=None, align=None, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Microsoft YaHei'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p

def add_quote(text, size=10.5):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Microsoft YaHei'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
    run.italic = True
    return p

def set_cell_text(cell, text, bold=False, size=9, color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.font.name = 'Microsoft YaHei'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color

def shade_cells(row, color_hex):
    for cell in row.cells:
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), color_hex)
        shading.set(qn('w:val'), 'clear')
        cell._tc.get_or_add_tcPr().append(shading)

# ════════════════════════════════════════
# 封面
# ════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
add_para('才链 TalenLink', bold=True, size=28, color=RGBColor(0x2B, 0x3A, 0x4E), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para('产教融合平台 · 宣传视频脚本', bold=True, size=16, color=RGBColor(0x44, 0x72, 0xC4), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=30)
add_para('━' * 40, size=10, color=RGBColor(0xCC, 0xCC, 0xCC), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=30)
add_para('视频时长：约 3 分钟', size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para('风格：写实 · 纪录片质感', size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para('适用场景：大创赛路演 · 项目答辩 · 宣传展示', size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
add_para('━' * 40, size=10, color=RGBColor(0xCC, 0xCC, 0xCC), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=30)

# 信息表
info_table = doc.add_table(rows=5, cols=2)
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
info_data = [
    ('项目名称', '才链 TalenLink'),
    ('核心定位', '缩短市场信号传导时间的人才培养链路平台'),
    ('视频时长', '约 3 分钟（170 秒）'),
    ('主角', '陈屿 · 大三 · 计算机科学与技术'),
    ('分镜数量', '13 个场景 · 5 幕'),
]
for i, (k, v) in enumerate(info_data):
    set_cell_text(info_table.rows[i].cells[0], k, bold=True, size=10, color=RGBColor(0x44, 0x72, 0xC4))
    set_cell_text(info_table.rows[i].cells[1], v, size=10)

doc.add_page_break()

# ════════════════════════════════════════
# 导演阐述
# ════════════════════════════════════════
add_heading_styled('导演阐述', 1)

add_heading_styled('导演五问', 2)
questions = [
    ('功能', '宣传片——但不是硬广，是一个人的真实经历。让评委在 3 分钟内理解：这个项目解决什么问题、怎么解决、对一个真实的人意味着什么。'),
    ('转折', '从"投了 37 份简历石沉大海的迷茫"→"我终于看见了市场在要什么"。'),
    ('视角', '陈屿，大三，计算机专业，普通本科。不是学霸，不是废柴，就是一个普通的、有点迷茫的大学生。观众代入他的身体。'),
    ('权力', '市场掌握信息差权力——学生看不见企业要什么，教师看不见市场变什么。平台把信号交到学生和教师手里。'),
    ('潜台词', '"不是我不够努力，是我一直看不见信号。"'),
]
for q, a in questions:
    p = doc.add_paragraph()
    run_q = p.add_run(f'{q}：')
    run_q.bold = True
    run_q.font.size = Pt(11)
    run_q.font.name = 'Microsoft YaHei'
    run_q.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run_a = p.add_run(a)
    run_a.font.size = Pt(11)
    run_a.font.name = 'Microsoft YaHei'
    run_a.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    p.paragraph_format.space_after = Pt(4)

add_heading_styled('一个意图', 2)
add_quote('让评委感受到——一个真实的人，因为一个平台，从看不见到看见了。', size=12)

add_heading_styled('视觉风格', 2)
style_items = [
    '写实摄影质感，自然肤色、自然光线、无滤镜美颜',
    '色调过渡：前期冷灰（宿舍/阴天）→ 中期暖色（图书馆/教室）→ 结尾金色（黄昏）',
    '主角身份全程一致：深色卫衣、黑色双肩包、体态从驼背逐渐挺直',
    '所有 UI/数据画面不使用可灵渲染——后期叠加真实平台录屏',
    '无文字、无 logo、无水印渲染',
]
for item in style_items:
    doc.add_paragraph(item, style='List Bullet')

add_heading_styled('声音设计', 2)
sound_items = [
    '旁白为主，对话为辅——旁白是陈屿的内心独白，对话是他和身边人的真实交流',
    '环境音写实：宿舍键盘声、校园自行车铃、图书馆翻书声、走廊回响',
    '配乐：低沉钢琴 + 轻电子氛围，节奏随情绪变化，结尾渐弱至安静',
]
for item in sound_items:
    doc.add_paragraph(item, style='List Bullet')

add_heading_styled('预算分配', 2)
budget_table = doc.add_table(rows=6, cols=4)
budget_table.style = 'Light Grid Accent 1'
budget_headers = ['段落', '主要消耗', '次要消耗', '有意节约']
for i, h in enumerate(budget_headers):
    set_cell_text(budget_table.rows[0].cells[i], h, bold=True, size=9)
shade_cells(budget_table.rows[0], '2B3A4E')
for run in budget_table.rows[0].cells[0].paragraphs[0].runs:
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
for run in budget_table.rows[0].cells[1].paragraphs[0].runs:
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
for run in budget_table.rows[0].cells[2].paragraphs[0].runs:
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
for run in budget_table.rows[0].cells[3].paragraphs[0].runs:
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

budget_data = [
    ('ACT 1 看不见', '情感共鸣', '场景密度', '产品细节'),
    ('ACT 2 看见了', '产品身份/UI', '人物反应', '动作幅度'),
    ('ACT 3 走进去了', '动作幅度', '场景密度', '面部特写'),
    ('ACT 4 看清了', '人物情感', '光线氛围', '产品信息'),
    ('ACT 5 看见未来', '氛围/意境', '人物身份', '对话'),
]
for i, (seg, main, sub, save) in enumerate(budget_data):
    set_cell_text(budget_table.rows[i+1].cells[0], seg, size=9)
    set_cell_text(budget_table.rows[i+1].cells[1], main, size=9)
    set_cell_text(budget_table.rows[i+1].cells[2], sub, size=9)
    set_cell_text(budget_table.rows[i+1].cells[3], save, size=9)

doc.add_page_break()

# ════════════════════════════════════════
# 角色设定
# ════════════════════════════════════════
add_heading_styled('角色设定', 1)

char_table = doc.add_table(rows=2, cols=6)
char_table.style = 'Light Grid Accent 1'
char_headers = ['姓名', '年龄', '身份', '专业', '外形', '性格']
for i, h in enumerate(char_headers):
    set_cell_text(char_table.rows[0].cells[i], h, bold=True, size=9)
shade_cells(char_table.rows[0], '2B3A4E')
for cell in char_table.rows[0].cells:
    for run in cell.paragraphs[0].runs:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

char_data = ['陈屿', '21', '大三本科生', '计算机科学与技术',
             '深色卫衣、黑色双肩包、\n微微驼背→逐渐挺直',
             '普通、内向、有点迷茫\n但不放弃']
for i, d in enumerate(char_data):
    set_cell_text(char_table.rows[1].cells[i], d, size=9)

add_para('')
add_para('体态变化线索：', bold=True, size=10)
posture = [
    'ACT 1-2：微微驼背，低头看手机，走路拖沓',
    'ACT 3：坐姿开始挺直，主动操作电脑',
    'ACT-4：肩膀打开，背挺直，双手插兜看向窗外',
    'ACT 5：步伐轻快，不再低头，看着前方走',
]
for item in posture:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ════════════════════════════════════════
# 完整分镜脚本
# ════════════════════════════════════════
add_heading_styled('完整分镜脚本', 1)

# ── Overview table ──
add_heading_styled('总览', 2)
overview_table = doc.add_table(rows=14, cols=5)
overview_table.style = 'Light Grid Accent 1'
ov_headers = ['编号', '时长', '场景', '核心', '对话类型']
for i, h in enumerate(ov_headers):
    set_cell_text(overview_table.rows[0].cells[i], h, bold=True, size=9)
shade_cells(overview_table.rows[0], '2B3A4E')
for cell in overview_table.rows[0].cells:
    for run in cell.paragraphs[0].runs:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

overview_data = [
    ('S01', '15s', '宿舍夜', '痛点：石沉大海', '旁白'),
    ('S02', '10s', '宿舍夜', '痛点：看不见市场', '旁白'),
    ('S03', '15s', '校园路', '同病相怜', '对话'),
    ('S04', '10s', '校园路', '孤独感', '自言自语'),
    ('S05', '15s', '图书馆', '发现平台', '旁白'),
    ('S06', '15s', '图书馆', '理解平台', '旁白'),
    ('S07', '10s', '图书馆', '释然', '旁白'),
    ('S08', '15s', '教室', '挑战广场', '对话'),
    ('S09', '15s', '教室', '真实编码', '旁白'),
    ('S10', '15s', '走廊窗边', '改变已发生', '对话'),
    ('S11', '15s', '走廊窗边', '连接感', '独白'),
    ('S12', '15s', '校园黄昏', '看见未来', '旁白'),
    ('S13', '10s', '校园黄昏', '结尾静帧', '无'),
]
for i, (num, dur, scene, core, dt) in enumerate(overview_data):
    row = overview_table.rows[i+1]
    set_cell_text(row.cells[0], num, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(row.cells[1], dur, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(row.cells[2], scene, size=9)
    set_cell_text(row.cells[3], core, size=9)
    set_cell_text(row.cells[4], dt, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

add_para('')
add_para('总时长：约 170 秒 ≈ 2 分 50 秒', bold=True, size=10)

doc.add_page_break()

# ════════════════════════════════════════
# 逐场景详细脚本
# ════════════════════════════════════════
add_heading_styled('逐场景详细脚本', 1)

# Scene data
scenes = [
    {
        'act': 'ACT 1 — 看不见',
        'num': 'S01',
        'title': '宿舍夜 · 石沉大海',
        'time': '15 秒',
        'shot': '中景 → 近景',
        'camera': '锁定，缓慢推进',
        'light': '笔记本屏幕蓝白光，宿舍其余全暗',
        'sound': '键盘敲击声、鼠标点击声、远处宿舍楼低沉人声',
        'action': '夜晚大学宿舍。陈屿坐在书桌前，笔记本屏幕的光打在脸上。桌上散落着几份打印的简历。他滑动鼠标，屏幕上是一个招聘页面——密密麻麻的列表在滚动。他停下，摘下眼镜揉眼睛，靠在椅背上。',
        'dialogue_type': '旁白',
        'dialogue': '"投了三十七份。石沉大海。"',
        'emotion': '疲惫、迷茫、压抑',
        'constraint': '写实摄影质感，自然皮肤纹理，无滤镜美颜。保持人物面部、发型、服装不变。无文字渲染。',
        'post': '叠加字幕（可选）："投了 37 份简历"',
        'prompt': '写实风格。夜晚大学四人间宿舍，一个 21 岁男生穿深色卫衣坐在书桌前，笔记本屏幕蓝白光照亮面部下半部分。桌上散落几份 A4 纸简历。他滑动鼠标滚轮，屏幕内容滚动，然后停下，摘下眼镜揉眼睛，靠在椅背上看向天花板。其余三个床铺无人，台灯关闭。镜头：中景锁定，缓慢推到近景。光线：屏幕蓝白主光，无其他光源。声音：键盘敲击、鼠标滚轮、远处宿舍楼低沉人声。约束：写实摄影质感，自然皮肤纹理，无滤镜美颜。保持人物面部、发型、服装不变。无文字渲染。',
    },
    {
        'act': '',
        'num': 'S02',
        'title': '宿舍夜 · 看不见市场',
        'time': '10 秒',
        'shot': '极近景',
        'camera': '锁定',
        'light': '屏幕蓝光',
        'sound': '鼠标点击声、一声叹气',
        'action': '屏幕极近景。模糊的蓝色光晕中，一行行文字在滚动。他的手指停在触摸板上，光标悬停不动。',
        'dialogue_type': '旁白',
        'dialogue': '"学了三年代码，不知道市场到底要什么样的人。"',
        'emotion': '困惑、无力',
        'constraint': '写实摄影质感。不渲染屏幕具体文字。保持手指和设备形态不变。',
        'post': '无',
        'prompt': '写实风格。极近景，笔记本电脑屏幕，蓝色光晕，一行行模糊的文字列表在滚动。一只手的手指停在触摸板上，光标不再移动。镜头：锁定极近景。光线：屏幕蓝白光。声音：鼠标点击声，一声轻微叹气。约束：写实摄影质感。不渲染屏幕具体文字。保持手指和设备形态不变。',
    },
    {
        'act': '',
        'num': 'S03',
        'title': '校园路 · 同病相怜',
        'time': '15 秒',
        'shot': '中远景 → 中景',
        'camera': '横向缓慢跟拍',
        'light': '阴天，灰白色漫射光',
        'sound': '校园环境音：自行车铃、远处篮球声、树叶沙沙',
        'action': '白天校园林荫道。陈屿背着黑色双肩包，微微驼背，耳机线垂在胸前。周围有同学说笑走过，他低着头看手机。一个同学从后面拍他肩膀。',
        'dialogue_type': '对话',
        'dialogue': '同学："屿哥，秋招怎么样了？"\n陈屿（没抬头，苦笑）："别提了。投了一堆，连面试都没几个。"\n同学："我也是。感觉学的跟人家要的完全对不上。"',
        'emotion': '无奈、共鸣',
        'constraint': '写实摄影质感，自然肤色，无美颜。保持主角面部和服装不变。无文字渲染。',
        'post': '无',
        'prompt': '写实风格。白天大学校园林荫道，阴天灰白色光线。一个 21 岁男生穿深色卫衣背黑色双肩包，微微驼背，戴耳机低头看手机走在路上。周围有其他学生走过。另一个男生从后面走上来拍他肩膀，他转头苦笑。镜头：中远景，横向缓慢跟拍。声音：自行车铃、远处篮球声、树叶沙沙、脚步声。约束：写实摄影质感，自然肤色，无美颜。保持主角面部和服装不变。无文字渲染。',
    },
    {
        'act': '',
        'num': 'S04',
        'title': '校园路 · 孤独感',
        'time': '10 秒',
        'shot': '中景',
        'camera': '锁定',
        'light': '阴天漫射',
        'sound': '校园环境音渐弱',
        'action': '两人并排走。陈屿摘下一只耳机，看向远处。同学拍了拍他背，拐向另一条路。陈屿一个人继续走，镜头停在原地，他的背影越来越小。',
        'dialogue_type': '自言自语',
        'dialogue': '"不是不努力……是真的看不见。"',
        'emotion': '孤独、低落',
        'constraint': '写实摄影质感。保持主角身份不变。无文字渲染。',
        'post': '无',
        'prompt': '写实风格。校园林荫道，阴天。两个男生并排走，一个拍了拍另一个的背然后拐向岔路。剩下的那个男生——深色卫衣、黑色双肩包、微微驼背——一个人继续向前走。镜头锁定中景，人物背影越来越小。声音：脚步声渐远，树叶沙沙，校园环境音渐弱。约束：写实摄影质感。保持主角身份不变。无文字渲染。',
    },
    {
        'act': 'ACT 2 — 看见了',
        'num': 'S05',
        'title': '图书馆 · 发现平台',
        'time': '15 秒',
        'shot': '中景 → 近景',
        'camera': '缓慢推进',
        'light': '暖色台灯 + 冷色窗户光混合',
        'sound': '图书馆安静环境音、翻书声、远处空调嗡鸣',
        'action': '图书馆靠窗位置。陈屿面前摊着课本和笔记本电脑。他习惯性地打开浏览器，手指悬在键盘上。手机屏幕亮了——一条消息推送。他拿起手机看了一眼，微微皱眉，然后在浏览器地址栏输入了什么。',
        'dialogue_type': '旁白',
        'dialogue': '"室友发了个链接，说是什么产教融合平台。我本来不想点的。"',
        'emotion': '好奇、将信将疑',
        'constraint': '写实摄影质感。保持人物身份不变。不渲染屏幕内容。无文字渲染。',
        'post': '叠加平台 UI 录屏（画中画/半透明叠加）；字幕："才链 TalenLink"',
        'prompt': '写实风格。大学图书馆靠窗位置，白天。一个 21 岁男生穿深色卫衣坐在桌前，面前摊着课本和笔记本电脑。暖色台灯从右侧照亮，冷色窗户光从左侧照入。他拿起手机看了一眼，微微皱眉，放下手机，在笔记本键盘上打字。镜头：中景缓慢推到近景。声音：图书馆安静环境音、翻书声、远处空调嗡鸣。约束：写实摄影质感。保持人物身份不变。不渲染屏幕内容。无文字渲染。',
    },
    {
        'act': '',
        'num': 'S06',
        'title': '图书馆 · 理解平台',
        'time': '15 秒',
        'shot': '近景（面部侧面）',
        'camera': '锁定',
        'light': '屏幕光映在面部',
        'sound': '鼠标点击、轻微页面加载音效',
        'action': '近景，陈屿的侧面。屏幕光映在脸上，表情在变化——从随意浏览，到微微前倾，到眼睛微微睁大。手指开始主动点击、滑动。嘴唇微动，像是在默读。',
        'dialogue_type': '旁白',
        'dialogue': '"它把十八万条真实岗位数据清洗过了，分成初级、中级、高级……还能看到你学的课跟市场要的技能差在哪。"',
        'emotion': '惊讶、理解、兴趣',
        'constraint': '写实摄影质感，自然面部细节。保持面部结构和发型不变。无文字渲染。',
        'post': '叠加平台功能截图；字幕："18.5 万真实岗位 · 3 级分层 · 职业画像"',
        'prompt': '写实风格。近景，21 岁男生侧面面部。笔记本屏幕光映在脸上，表情从随意变成微微前倾、眼睛微睁大。手指在触摸板上主动滑动和点击。嘴唇微动像在默读。镜头：锁定近景。光线：屏幕蓝白光映面部。声音：鼠标点击、轻微页面音效。约束：写实摄影质感，自然面部细节。保持面部结构和发型不变。无文字渲染。',
    },
    {
        'act': '',
        'num': 'S07',
        'title': '图书馆 · 释然',
        'time': '10 秒',
        'shot': '极近景 → 近景',
        'camera': '锁定',
        'light': '屏幕光 + 台灯暖光',
        'sound': '点击声、一声轻笑',
        'action': '极近景，手指点击屏幕上一个按钮区域。手指停住。然后——他笑了。嘴角微微上扬，像是终于看到了什么。',
        'dialogue_type': '旁白',
        'dialogue': '"原来不是我不够好……是我一直看不见信号。"',
        'emotion': '释然、如释重负',
        'constraint': '写实摄影质感。保持手部和面部身份不变。不渲染屏幕具体内容。无文字渲染。',
        'post': '无',
        'prompt': '写实风格。极近景，一只手的手指点击笔记本屏幕上一个按钮区域。手指停住。然后切到近景，男生嘴角微微上扬，像是释然的笑。不是大笑，是安静的、轻微的微笑。镜头：极近景锁定。光线：屏幕光 + 台灯暖光。声音：点击声、一声轻笑。约束：写实摄影质感。保持手部和面部身份不变。不渲染屏幕具体内容。无文字渲染。',
    },
    {
        'act': 'ACT 3 — 走进去了',
        'num': 'S08',
        'title': '教室 · 挑战广场',
        'time': '15 秒',
        'shot': '远景 → 中景',
        'camera': '缓慢推进',
        'light': '教室日光灯 + 投影仪光',
        'sound': '教室环境音、键盘声、低声讨论',
        'action': '学校计算机教室。陈屿坐在中间位置，屏幕上是他之前没见过的界面。周围有其他同学也在操作。他点进一个企业发布的技术挑战，看了看要求，回头看了一眼旁边的同学，两人交换了一个眼神。',
        'dialogue_type': '对话',
        'dialogue': '陈屿（低声）："这个是企业直接出的真实需求？不是那种假的课程设计？"\n同学（点头，压低声音）："嗯，做完企业直接评审。通过了还有证书。"',
        'emotion': '好奇、认真',
        'constraint': '写实摄影质感。保持主角身份和服装不变。不渲染屏幕具体内容。无文字渲染。',
        'post': '叠加挑战广场 UI；字幕："揭榜挂帅 · 以战代练"',
        'prompt': '写实风格。大学计算机教室，白天，日光灯和投影仪光。一个 21 岁男生穿深色卫衣坐在电脑前，屏幕上是抽象的界面。周围有其他同学在操作电脑。他看了一眼屏幕，转头和旁边同学低声交谈，同学点头回应。镜头：远景缓慢推到中景。声音：教室环境音、键盘声、低声讨论。约束：写实摄影质感。保持主角身份和服装不变。不渲染屏幕具体内容。无文字渲染。',
    },
    {
        'act': '',
        'num': 'S09',
        'title': '教室 · 真实编码',
        'time': '15 秒',
        'shot': '中近景',
        'camera': '缓慢推进',
        'light': '屏幕光照脸',
        'sound': '键盘敲击、鼠标点击、偶尔低声讨论',
        'action': '时间推移。陈屿在认真写代码，屏幕上是代码编辑器的抽象光效。他停下来想了一下，然后快速敲了一段。旁边同学探过头看了一眼他的屏幕，竖了个大拇指。陈屿没抬头，嘴角微动。',
        'dialogue_type': '旁白',
        'dialogue': '"第一次写代码不是为了交作业，是为了一个真实的问题。那种感觉……不一样。"',
        'emotion': '专注、充实、成就感',
        'constraint': '写实摄影质感。保持身份不变。不渲染屏幕具体代码。无文字渲染。',
        'post': '无',
        'prompt': '写实风格。大学计算机教室。21 岁男生穿深色卫衣在电脑前认真打字，屏幕上是代码编辑器的抽象蓝绿色光效。他停下来思考，然后快速打字。旁边同学探头看了一眼屏幕，竖起大拇指。他没抬头，嘴角微微一动。镜头：中近景缓慢推进。光线：屏幕光照脸。声音：键盘敲击、鼠标点击、低声讨论。约束：写实摄影质感。保持身份不变。不渲染屏幕具体代码。无文字渲染。',
    },
    {
        'act': 'ACT 4 — 看清了',
        'num': 'S10',
        'title': '走廊窗边 · 改变已发生',
        'time': '15 秒',
        'shot': '中景 → 中近景',
        'camera': '缓慢推进',
        'light': '走廊窗户自然光，温暖',
        'sound': '走廊脚步声回响、远处操场广播',
        'action': '教学楼走廊，下午阳光斜射进来。陈屿靠在窗边，手里拿着手机。他的站姿变了——不再是驼背，而是肩膀打开，背挺直。他抬头看向窗外操场，阳光打在脸上。',
        'dialogue_type': '对话（对镜头外的人）',
        'dialogue': '"你知道最离谱的是什么吗？我之前投的那些简历，岗位要求我都没看懂过。不是写不出来，是根本不知道人家要什么。"\n（顿了一下）\n"现在知道了。"',
        'emotion': '平静、自信、释然',
        'constraint': '写实摄影质感，阳光自然色温。保持身份不变。无文字渲染。',
        'post': '无',
        'prompt': '写实风格。大学教学楼走廊，下午阳光从窗户斜射进来。21 岁男生穿深色卫衣靠在窗边看手机，站姿挺直，肩膀打开——不再是驼背。他抬头看向窗外操场，阳光照在脸上。镜头：中景缓慢推到中近景。光线：温暖自然窗户光。声音：走廊脚步声回响、远处操场广播。约束：写实摄影质感，阳光自然色温。保持身份不变。无文字渲染。',
    },
    {
        'act': '',
        'num': 'S11',
        'title': '走廊窗边 · 连接感',
        'time': '15 秒',
        'shot': '中近景',
        'camera': '锁定',
        'light': '侧光，阳光',
        'sound': '走廊环境音、远处有人在笑',
        'action': '继续。陈屿收起手机，双手插兜，看向窗外。风吹进来，头发微微动。他深吸一口气，像是做了个决定。',
        'dialogue_type': '独白',
        'dialogue': '"这个平台……让我第一次觉得，学校学的东西跟外面的世界是连着的。"',
        'emotion': '平静、坚定',
        'constraint': '写实摄影质感。保持身份不变。无文字渲染。',
        'post': '字幕："产教融合 · 数据驱动"',
        'prompt': '写实风格。教学楼走廊窗边，下午阳光。21 岁男生穿深色卫衣收起手机，双手插兜看向窗外。微风吹动头发。他深吸一口气，表情平静但有决心。镜头：中近景锁定。光线：侧面阳光，暖色。声音：走廊环境音、远处笑声、微风。约束：写实摄影质感。保持身份不变。无文字渲染。',
    },
    {
        'act': 'ACT 5 — 看见未来',
        'num': 'S12',
        'title': '校园黄昏 · 看见未来',
        'time': '15 秒',
        'shot': '中远景 → 远景',
        'camera': '缓慢后退',
        'light': '黄昏金色逆光',
        'sound': '校园傍晚环境音：远处吉他声、自行车铃、鸟鸣',
        'action': '黄昏，校园主路。陈屿背着包走在路上，步伐比之前轻快。周围是放学的人流。他没有低头看手机，而是看着前方。夕阳在他身后，逆光勾出轮廓。他越走越远，融入人流。',
        'dialogue_type': '旁白',
        'dialogue': '"看见了信号，就知道该往哪走了。"',
        'emotion': '平静、坚定、希望',
        'constraint': '写实摄影质感，黄昏自然色温。保持身份不变。无文字渲染。',
        'post': '叠加项目名 + 定位语',
        'prompt': '写实风格。大学校园主路，黄昏金色逆光。21 岁男生穿深色卫衣背黑色双肩包走在路上，步伐轻快，不再驼背，看着前方。周围有放学的学生人流。夕阳在他身后，逆光勾出身体轮廓。他越走越远，融入人群。镜头：中远景缓慢后退到远景。声音：远处吉他声、自行车铃、鸟鸣、脚步声。约束：写实摄影质感，黄昏自然色温。保持身份不变。无文字渲染。',
    },
    {
        'act': '',
        'num': 'S13',
        'title': '校园黄昏 · 结尾静帧',
        'time': '10 秒',
        'shot': '远景',
        'camera': '缓慢后退至静止',
        'light': '黄昏',
        'sound': '环境音渐弱至安静',
        'action': '陈屿的背影在校园路上越来越小，画面逐渐变暗。最后一帧：他的剪影在夕阳中。',
        'dialogue_type': '无',
        'dialogue': '（静默）',
        'emotion': '余韵、希望',
        'constraint': '写实摄影质感。保持身份不变。无文字渲染。',
        'post': '渐显项目信息："才链 TalenLink · 缩短市场信号传导时间"',
        'prompt': '写实风格。校园黄昏远景。一个穿深色卫衣背黑色双肩包的男生背影在校园路上越走越远，画面缓慢变暗。最后定格在人物剪影和夕阳。镜头：远景缓慢后退至静止。光线：黄昏逆光。声音：环境音渐弱至安静。约束：写实摄影质感。保持身份不变。无文字渲染。',
    },
]

current_act = ''
for s in scenes:
    # Act header
    if s['act'] and s['act'] != current_act:
        current_act = s['act']
        add_heading_styled(current_act, 2, color=RGBColor(0x44, 0x72, 0xC4))

    # Scene header
    add_heading_styled(f"{s['num']} · {s['title']}", 3)

    # Info table
    t = doc.add_table(rows=6, cols=2)
    t.style = 'Light Grid Accent 1'
    info_rows = [
        ('时长', s['time']),
        ('景别', s['shot']),
        ('镜头', s['camera']),
        ('光线', s['light']),
        ('声音', s['sound']),
        ('对话类型', s['dialogue_type']),
    ]
    for i, (k, v) in enumerate(info_rows):
        set_cell_text(t.rows[i].cells[0], k, bold=True, size=9)
        set_cell_text(t.rows[i].cells[1], v, size=9)

    add_para('')

    # Action
    add_para('画面描述', bold=True, size=10, color=RGBColor(0x2B, 0x3A, 0x4E))
    add_para(s['action'], size=10)

    # Dialogue
    add_para('台词', bold=True, size=10, color=RGBColor(0x44, 0x72, 0xC4))
    for line in s['dialogue'].split('\n'):
        add_quote(line.strip(), size=10)

    # Emotion
    add_para(f"情感基调：{s['emotion']}", italic=True, size=10, color=RGBColor(0x66, 0x66, 0x66))

    # Constraint
    add_para(f"约束：{s['constraint']}", size=9, color=RGBColor(0x88, 0x88, 0x88))

    # Post
    if s['post'] != '无':
        add_para(f"后期：{s['post']}", size=9, color=RGBColor(0x44, 0x72, 0xC4))

    # Prompt
    add_para('可灵 AI Prompt', bold=True, size=10, color=RGBColor(0x2B, 0x3A, 0x4E))
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(s['prompt'])
    run.font.name = 'Consolas'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    # Light gray background for prompt
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'F5F5F5')
    shading.set(qn('w:val'), 'clear')
    p.runs[0].element.rPr.append(shading)

    # Separator
    add_para('━' * 60, size=8, color=RGBColor(0xDD, 0xDD, 0xDD), space_after=12)

doc.add_page_break()

# ════════════════════════════════════════
# 后期制作指南
# ════════════════════════════════════════
add_heading_styled('后期制作指南', 1)

add_heading_styled('UI 叠加方案', 2)
ui_items = [
    'S05-S06：将平台真实录屏以画中画或半透明叠加的方式嵌入画面，与可灵生成的人物画面合成',
    'S08：挑战广场界面叠加在教室电脑屏幕上',
    'S09：代码编辑器界面可选择性叠加',
    '所有 UI 内容使用平台真实录屏，不要用可灵渲染——文字/logo 会变形',
]
for item in ui_items:
    doc.add_paragraph(item, style='List Bullet')

add_heading_styled('字幕方案', 2)
sub_items = [
    '旁白字幕：白色，底部居中，微软雅黑，24px，半透明黑底',
    '数据标注字幕（S06）：#4472C4 蓝色，右下角，小字',
    '项目信息（S13）：居中，渐显动画，大字',
]
for item in sub_items:
    doc.add_paragraph(item, style='List Bullet')

add_heading_styled('配乐建议', 2)
music_items = [
    'ACT 1-2（S01-S07）：低沉钢琴单音 + 轻电子氛围，节奏慢，情绪压抑→逐渐明亮',
    'ACT 3（S08-S09）：节奏加快，加入轻微鼓点，代表行动和参与',
    'ACT 4-5（S10-S13）：钢琴旋律变温暖，电子音渐弱，最后一段只剩钢琴，渐弱至安静',
    '整体音量：旁白时配乐降至 -12dB，无人声时恢复',
]
for item in music_items:
    doc.add_paragraph(item, style='List Bullet')

add_heading_styled('色调参考', 2)
color_items = [
    'ACT 1：冷灰调，低饱和，偏蓝灰（#6B7B8D）',
    'ACT 2：逐渐加入暖色，台灯暖光开始出现',
    'ACT 3：中性色，教室日光灯标准色温',
    'ACT 4：暖色为主，下午阳光（#E8D5B7）',
    'ACT 5：金色黄昏（#D4A574），逆光，高光过曝有意为之',
]
for item in color_items:
    doc.add_paragraph(item, style='List Bullet')

add_heading_styled('剪辑节奏', 2)
edit_items = [
    'S01-S04（痛点）：慢切，每个镜头停留时间长，让观众感受压抑',
    'S05-S07（发现）：切速逐渐加快，代表信息涌入',
    'S08-S09（行动）：节奏稳定，代表专注',
    'S10-S13（成长）：回到慢切，但这次是从容的慢，不是压抑的慢',
]
for item in edit_items:
    doc.add_paragraph(item, style='List Bullet')

# ════════════════════════════════════════
# 保存
# ════════════════════════════════════════
output_path = r'C:\Users\YY\kling-prompt-engineering\才链_TalenLink_宣传视频脚本.docx'
doc.save(output_path)
print(f'文档已保存: {output_path}')
